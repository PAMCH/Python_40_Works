import docx # 워드 파일 작성을 위한 모듈
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document(r'수료증양식.docx') # 양식 파일을 불러옴
#기본 폰트와 글씨 크기 지정
style = doc.styles['Normal']
style.font.name = '나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
style.font.size = docx.shared.Pt(12)
#문단 생성 및 문자 입력 후 크기 및 폰트 지정
para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('제 2020-0001호\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
#본문
para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('수 료 증')
run.font.name = '나눔고딕'
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run(' 성 명: 장다인\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 생 년 월 일 : 2017.12.12\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 교 육 과 정: 파이썬 과 40개의 작품들\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 교 육 날 짜: 2021.08.05~2021.09.09\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('위 사람은 파이썬과 40개의 작품들 교육 과정을\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run = para.add_run(' 이수하였으므로 이 증서를 수여 합니다.')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('2021.09.19')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para = doc.add_paragraph()
run = para.add_run('\n\n')
run = para.add_run('파이썬교육기관장')
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'수료증결과.docx') # 양식파일에 글을 더한 내용을 저장